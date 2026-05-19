from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_blank_resume_docx import (
    ACCENT,
    MUTED,
    add_labeled_grid,
    add_paragraph_text,
    add_section_heading,
    add_single_column_grid,
    style_run,
)


OUTPUT_PATH = Path(__file__).resolve().parent / "example_resume_sample.docx"
SEED_PATH = Path(__file__).resolve().parent.parent / "seed" / "my_info_data.json"


def build_document() -> None:
    profile = json.loads(SEED_PATH.read_text(encoding="utf-8"))

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("이력서")
    style_run(title_run, size=20, bold=True, color=ACCENT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    add_paragraph_text(
        subtitle,
        f"{profile['이름']} | {profile['희망직무']}",
        size=10,
        color=MUTED,
    )

    add_section_heading(document, "기본 인적사항")
    add_labeled_grid(
        document,
        [
            ["이름", profile["이름"], "생년월일", profile["생년월일"]],
            ["나이", profile["나이"], "이메일", profile["이메일"]],
            ["전화번호", profile["전화번호"], "주소", profile["주소"]],
        ],
        widths=(Cm(2.4), Cm(6.0), Cm(2.6), Cm(6.0)),
    )

    add_section_heading(document, "지원 정보")
    add_single_column_grid(
        document,
        [
            ("희망직무", profile["희망직무"]),
            ("기술스택", ", ".join(profile["기술스택"])),
            ("자격증", ", ".join(profile["자격증"])),
            ("취미", ", ".join(profile["취미"])),
        ],
        label_width=Cm(4.2),
        value_width=Cm(12.2),
    )

    add_section_heading(document, "학력")
    add_labeled_grid(
        document,
        [
            ["대학교", profile["학력"]["대학교"], "전공", profile["학력"]["전공"]],
            ["학점", profile["학력"]["학점"], "졸업년도", profile["학력"]["졸업년도"]],
            ["경력", profile["경력"], "추가메모", profile.get("추가메모", "")],
        ],
        widths=(Cm(2.8), Cm(5.6), Cm(3.0), Cm(5.6)),
    )

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    add_section_heading(document, "자기소개")
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    add_paragraph_text(
        intro,
        f"{profile['이름']}은(는) {profile['학력']['대학교']} {profile['학력']['전공']} 재학 중이며, "
        f"{profile['희망직무']} 직무를 목표로 {', '.join(profile['기술스택'][:4])} 중심의 개발 역량을 준비하고 있습니다.",
        size=10,
    )

    add_section_heading(document, "직무 관련 역량")
    strengths = [
        f"주요 기술: {', '.join(profile['기술스택'])}",
        f"보유 자격증: {', '.join(profile['자격증'])}",
        "로컬 AI 기반 자동입력, 게임 클라이언트 개발, 일반 소프트웨어 구현 과제에 활용 가능한 기술 기반 보유",
    ]
    for item in strengths:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.4)
        add_paragraph_text(paragraph, f"- {item}", size=10)

    add_section_heading(document, "예시 자기소개 문장")
    sample = document.add_paragraph()
    sample.paragraph_format.space_after = Pt(10)
    add_paragraph_text(
        sample,
        "사용자 정보 자동화와 실용적인 개발 경험을 결합해, 반복적인 문서 작성 업무를 효율화할 수 있는 개발자가 되고자 합니다.",
        size=10,
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
