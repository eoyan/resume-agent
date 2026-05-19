from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parent / "autofill_resume_template.docx"

ACCENT = RGBColor(34, 94, 72)
ACCENT_SOFT = "E6F2EC"
LINE = "CBD8D1"
TEXT = RGBColor(27, 34, 31)
MUTED = RGBColor(96, 109, 103)
PLACEHOLDER = RGBColor(170, 179, 175)
FONT_NAME = "Malgun Gothic"


def set_font(run, *, size: int, bold: bool = False, color: RGBColor = TEXT) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_text(paragraph, text: str, *, size: int = 10, bold: bool = False, color: RGBColor = TEXT):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table_cell(cell, *, shaded: bool = False) -> None:
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if shaded:
        set_cell_shading(cell, "F5F8F6")


def add_section_band(document: Document, title: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    style_table_cell(cell)
    set_cell_shading(cell, ACCENT_SOFT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, title, size=11, bold=True, color=ACCENT)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_four_col_table(document: Document, rows: list[list[str]], widths: tuple[Cm, Cm, Cm, Cm]) -> None:
    table = document.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for column, width in zip(table.columns, widths):
        column.width = width
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            style_table_cell(cell, shaded=(col_index % 2 == 0))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, size=10, bold=(col_index % 2 == 0), color=MUTED if col_index % 2 == 0 else TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_two_col_table(document: Document, rows: list[tuple[str, str]], widths: tuple[Cm, Cm]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = widths[0]
    table.columns[1].width = widths[1]
    for row_index, (label, value) in enumerate(rows):
        left = table.cell(row_index, 0)
        right = table.cell(row_index, 1)
        style_table_cell(left, shaded=True)
        style_table_cell(right)
        add_text(left.paragraphs[0], label, size=10, bold=True, color=MUTED)
        add_text(right.paragraphs[0], value, size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_large_field_table(document: Document, rows: list[tuple[str, int]], *, label_width: Cm, value_width: Cm) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = label_width
    table.columns[1].width = value_width
    for row_index, (label, line_count) in enumerate(rows):
        left = table.cell(row_index, 0)
        right = table.cell(row_index, 1)
        style_table_cell(left, shaded=True)
        style_table_cell(right)
        add_text(left.paragraphs[0], label, size=10, bold=True, color=MUTED)
        first = right.paragraphs[0]
        add_text(first, "", size=10)
        for _ in range(max(1, line_count) - 1):
            paragraph = right.add_paragraph()
            add_text(paragraph, "", size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_prompt_block(document: Document, number: int, question: str, lines: int) -> None:
    question_p = document.add_paragraph()
    question_p.paragraph_format.space_after = Pt(5)
    add_text(question_p, f"{number}. {question}", size=10, bold=True)

    box = document.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Cm(16.2)
    cell = box.cell(0, 0)
    style_table_cell(cell)
    set_cell_margins(cell, top=140, start=130, bottom=140, end=130)
    placeholder = cell.paragraphs[0]
    add_text(placeholder, "", size=10)
    for _ in range(lines - 1):
        paragraph = cell.add_paragraph()
        add_text(paragraph, "", size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_hint_box(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    style_table_cell(cell)
    set_cell_shading(cell, "FAFCFB")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, text, size=9, color=MUTED)
    document.add_paragraph().paragraph_format.space_after = Pt(6)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, "자동입력 테스트용 이력서 양식", size=18, bold=True, color=ACCENT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    add_text(subtitle, "DOCX 업로드 테스트를 위한 구조화된 빈 양식", size=10, color=MUTED)

    add_hint_box(
        document,
        "이 템플릿은 PDF 좌표 문제를 피하기 위해 DOCX 기반 자동작성 테스트용으로 설계했습니다. 라벨이 명확하고 서술형 영역이 넓어 결과 확인과 후편집이 쉽습니다.",
    )

    add_section_band(document, "기본 정보")
    add_four_col_table(
        document,
        [
            ["이름", "", "생년월일", ""],
            ["전화번호", "", "이메일", ""],
            ["희망직무", "", "지원 회사명", ""],
            ["주소", "", "포트폴리오 URL", ""],
        ],
        widths=(Cm(2.8), Cm(5.5), Cm(3.0), Cm(5.0)),
    )

    add_section_band(document, "학력")
    add_four_col_table(
        document,
        [
            ["학교명", "", "전공", ""],
            ["부전공", "", "학점", ""],
            ["재학/졸업 상태", "", "졸업(예정)년도", ""],
        ],
        widths=(Cm(3.2), Cm(5.0), Cm(3.4), Cm(4.7)),
    )

    add_section_band(document, "경력 및 프로젝트")
    add_large_field_table(
        document,
        [
            ("총 경력", 1),
            ("최근 근무 회사", 1),
            ("담당 업무", 3),
            ("주요 성과", 3),
            ("주요 프로젝트", 4),
        ],
        label_width=Cm(4.0),
        value_width=Cm(12.4),
    )

    add_section_band(document, "기술 및 자격")
    add_two_col_table(
        document,
        [
            ("보유 기술스택", ""),
            ("사용 가능 언어", ""),
            ("협업 툴 경험", ""),
            ("보유 자격증", ""),
            ("수료 교육", ""),
        ],
        widths=(Cm(4.0), Cm(12.4)),
    )

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    add_section_band(document, "자기소개 및 서술형 문항")
    add_prompt_block(document, 1, "본인을 간단히 소개해 주세요.", 5)
    add_prompt_block(document, 2, "지원 직무와 관련된 강점을 작성해 주세요.", 5)
    add_prompt_block(document, 3, "가장 자신 있는 기술 또는 프로젝트 경험을 작성해 주세요.", 5)
    add_prompt_block(document, 4, "지원 동기를 작성해 주세요.", 4)
    add_prompt_block(document, 5, "입사 후 포부를 작성해 주세요.", 4)

    add_section_band(document, "추가 확인 항목")
    add_two_col_table(
        document,
        [
            ("Github / 블로그", ""),
            ("병역 사항", ""),
            ("수상 경력", ""),
            ("어학 점수", ""),
            ("기타 참고사항", ""),
        ],
        widths=(Cm(4.0), Cm(12.4)),
    )

    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer_p, "Digital Persona Autofill DOCX Template", size=8, color=MUTED)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
