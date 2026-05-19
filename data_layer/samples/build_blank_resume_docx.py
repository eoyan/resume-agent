from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(__file__).resolve().parent / "blank_resume_sample.docx"


ACCENT = RGBColor(21, 107, 82)
ACCENT_SOFT = "DFF2EA"
LINE_COLOR = "C9D6D0"
TEXT = RGBColor(24, 32, 29)
MUTED = RGBColor(102, 115, 109)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = LINE_COLOR, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, *, size: int, bold: bool = False, color: RGBColor = TEXT) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_paragraph_text(paragraph, text: str, *, size: int = 10, bold: bool = False, color: RGBColor = TEXT):
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    return run


def add_section_heading(document: Document, title: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_SOFT)
    set_cell_border(cell, color="B5C7BF", size="8")
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title)
    style_run(run, size=11, bold=True, color=ACCENT)
    document.add_paragraph()


def add_labeled_grid(document: Document, rows: list[list[str]], *, widths: tuple[Cm, Cm, Cm, Cm]) -> None:
    table = document.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col, width in zip(table.columns, widths):
        col.width = width

    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if col_idx % 2 == 0:
                set_cell_shading(cell, "F4F7F5")
                add_paragraph_text(paragraph, value, size=10, bold=True, color=MUTED)
            else:
                add_paragraph_text(paragraph, value, size=10)
    document.add_paragraph()


def add_single_column_grid(document: Document, rows: list[tuple[str, str]], *, label_width: Cm, value_width: Cm) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = label_width
    table.columns[1].width = value_width

    for row_idx, (label, value) in enumerate(rows):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)
        for cell in (label_cell, value_cell):
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(label_cell, "F4F7F5")
        add_paragraph_text(label_cell.paragraphs[0], label, size=10, bold=True, color=MUTED)
        add_paragraph_text(value_cell.paragraphs[0], value, size=10)
    document.add_paragraph()


def add_blank_answer_block(document: Document, question_number: int, title: str, lines: int = 4) -> None:
    question = document.add_paragraph()
    question.paragraph_format.space_after = Pt(5)
    add_paragraph_text(question, f"{question_number}. {title}", size=10, bold=True)

    box = document.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = True
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C7D1CC", size="6")
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    for index in range(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("_" * 110)
        style_run(run, size=9, color=RGBColor(185, 193, 189))
    document.add_paragraph()


def add_bullet_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Cm(0.4)
    run = paragraph.add_run(f"- {text}")
    style_run(run, size=10, color=MUTED)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("이력서 / 입사지원서")
    style_run(title_run, size=18, bold=True, color=ACCENT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    add_paragraph_text(
        subtitle,
        "디지털 페르소나 자동입력 테스트용 빈 문서 샘플",
        size=10,
        color=MUTED,
    )

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(14)
    add_paragraph_text(
        note,
        "각 항목을 비워둔 상태로 작성했으며, 업로드 후 로컬 AI가 자동으로 초안을 채우는 테스트에 사용할 수 있습니다.",
        size=9,
        color=MUTED,
    )

    add_section_heading(document, "기본 인적사항")
    add_labeled_grid(
        document,
        [
            ["이름", "", "생년월일", ""],
            ["나이", "", "이메일", ""],
            ["전화번호", "", "주소", ""],
        ],
        widths=(Cm(2.4), Cm(6.0), Cm(2.6), Cm(6.0)),
    )

    add_section_heading(document, "지원 정보")
    add_single_column_grid(
        document,
        [
            ("지원 회사명", ""),
            ("지원 부서", ""),
            ("지원 직무", ""),
            ("희망 연봉", ""),
            ("입사 가능 시기", ""),
        ],
        label_width=Cm(4.2),
        value_width=Cm(12.2),
    )

    add_section_heading(document, "학력")
    add_labeled_grid(
        document,
        [
            ["학교명", "", "전공", ""],
            ["부전공", "", "학점", ""],
            ["재학 상태", "", "졸업(예정)년도", ""],
        ],
        widths=(Cm(2.8), Cm(5.6), Cm(3.0), Cm(5.6)),
    )

    add_section_heading(document, "경력")
    add_single_column_grid(
        document,
        [
            ("총 경력", ""),
            ("최근 근무 회사", ""),
            ("담당 업무", ""),
            ("주요 성과", ""),
        ],
        label_width=Cm(4.2),
        value_width=Cm(12.2),
    )

    add_section_heading(document, "기술 및 역량")
    add_single_column_grid(
        document,
        [
            ("보유 기술스택", ""),
            ("사용 가능 언어", ""),
            ("주요 프로젝트", ""),
            ("협업 툴 경험", ""),
        ],
        label_width=Cm(4.2),
        value_width=Cm(12.2),
    )

    add_section_heading(document, "자격증 및 교육")
    add_single_column_grid(
        document,
        [
            ("보유 자격증", ""),
            ("수료 교육", ""),
        ],
        label_width=Cm(4.2),
        value_width=Cm(12.2),
    )

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    add_section_heading(document, "자기소개 및 서술형 문항")
    add_blank_answer_block(document, 1, "본인을 간단히 소개해 주세요.", lines=4)
    add_blank_answer_block(document, 2, "지원 직무와 관련된 강점을 작성해 주세요.", lines=5)
    add_blank_answer_block(document, 3, "가장 자신 있는 기술 또는 프로젝트 경험을 작성해 주세요.", lines=5)
    add_blank_answer_block(document, 4, "지원 동기를 작성해 주세요.", lines=4)
    add_blank_answer_block(document, 5, "입사 후 포부를 작성해 주세요.", lines=4)

    add_section_heading(document, "추가 확인 필요")
    for item in [
        "포트폴리오 URL",
        "Github / 블로그",
        "병역 사항",
        "수상 경력",
        "어학 점수",
        "기타 참고사항",
    ]:
        add_bullet_line(document, f"{item}: ________________________________________________")

    footer = document.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_text(
        footer_paragraph,
        "Local Persona Autofill Test Template",
        size=8,
        color=MUTED,
    )

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
