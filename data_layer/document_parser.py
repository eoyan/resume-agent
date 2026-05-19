from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import List


BUNDLED_SITE_PACKAGES = Path(
    "/Users/ny/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/lib/python3.12/site-packages"
)

if BUNDLED_SITE_PACKAGES.exists():
    sys.path.append(str(BUNDLED_SITE_PACKAGES))

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None


class DocumentParseError(RuntimeError):
    pass


class DocumentParser:
    supported_extensions = {".txt", ".md", ".json", ".docx", ".pdf"}

    def extract_text(self, file_path: Path) -> str:
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in self.supported_extensions:
            return self._clean_text(path.read_text(encoding="utf-8", errors="ignore"))
        if extension in {".txt", ".md"}:
            return self._clean_text(path.read_text(encoding="utf-8", errors="ignore"))
        if extension == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            return json.dumps(data, ensure_ascii=False, indent=2)
        if extension == ".docx":
            return self._extract_docx(path)
        if extension == ".pdf":
            return self._extract_pdf(path)
        return self._clean_text(path.read_text(encoding="utf-8", errors="ignore"))

    def preview(self, text: str, limit: int = 2400) -> str:
        lines: List[str] = []
        total = 0
        previous_blank = False

        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                if previous_blank:
                    continue
                snippet = ""
                previous_blank = True
            else:
                snippet = line
                previous_blank = False

            projected = total + len(snippet) + 1
            if projected > limit:
                remaining = max(0, limit - total)
                if remaining > 0 and snippet:
                    lines.append(snippet[:remaining].rstrip())
                lines.append("[미리보기 생략]")
                break

            lines.append(snippet)
            total = projected

        return "\n".join(lines).strip()

    def _extract_docx(self, file_path: Path) -> str:
        if Document is None:
            raise DocumentParseError("python-docx를 사용할 수 없습니다.")

        document = Document(str(file_path))
        blocks: List[str] = []

        for paragraph in document.paragraphs:
            text = self._clean_line(paragraph.text)
            if text:
                blocks.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [self._clean_line(cell.text) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    blocks.append(" | ".join(cells))

        return self._clean_text("\n".join(blocks).strip())

    def _extract_pdf(self, file_path: Path) -> str:
        if PdfReader is None:
            raise DocumentParseError("pypdf를 사용할 수 없습니다.")

        reader = PdfReader(str(file_path))
        pages: List[str] = []

        for index, page in enumerate(reader.pages, start=1):
            page_text = self._extract_pdf_page_text(page)
            page_text = self._clean_text(page_text)
            if page_text:
                pages.append(f"[PAGE {index}]\n{page_text}")

        if not pages:
            raise DocumentParseError("PDF에서 텍스트를 추출하지 못했습니다.")
        return "\n\n".join(pages)

    def _extract_pdf_page_text(self, page) -> str:
        candidates = []
        try:
            candidates.append(page.extract_text(extraction_mode="layout") or "")
        except TypeError:
            pass
        except Exception:
            pass

        try:
            candidates.append(page.extract_text() or "")
        except Exception:
            pass

        best = ""
        for candidate in candidates:
            if candidate and len(candidate.strip()) > len(best.strip()):
                best = candidate
        return best

    def _clean_text(self, text: str) -> str:
        normalized_lines: List[str] = []
        previous_blank = False

        for raw_line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            line = self._clean_line(raw_line)
            if not line:
                if previous_blank:
                    continue
                normalized_lines.append("")
                previous_blank = True
                continue

            normalized_lines.append(line)
            previous_blank = False

        return "\n".join(normalized_lines).strip()

    def _clean_line(self, line: str) -> str:
        text = line.replace(" ", " ").strip()
        text = re.sub(r"[ 	]+", " ", text)
        return text
