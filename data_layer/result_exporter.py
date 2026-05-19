from __future__ import annotations

import re
import sys
import uuid
from datetime import datetime
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

BUNDLED_SITE_PACKAGES = Path(
    "/Users/ny/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/lib/python3.12/site-packages"
)

if BUNDLED_SITE_PACKAGES.exists():
    sys.path.append(str(BUNDLED_SITE_PACKAGES))

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None

try:
    from pypdf import PdfReader, PdfWriter  # type: ignore
    from pypdf.generic import (  # type: ignore
        ArrayObject,
        BooleanObject,
        DictionaryObject,
        NameObject,
        TextStringObject,
    )
except Exception:
    PdfReader = None
    PdfWriter = None
    ArrayObject = None
    BooleanObject = None
    DictionaryObject = None
    NameObject = None
    TextStringObject = None

try:
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
except Exception:
    pdfmetrics = None
    TTFont = None
    canvas = None


class ResultExportError(RuntimeError):
    pass


class ResultExporter:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.pdf_font_name = self._configure_pdf_font()

    def export_autofill_result(
        self,
        *,
        original_name: str,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
        source_document_path: Optional[Path] = None,
    ) -> List[Dict[str, str]]:
        downloads: List[Dict[str, str]] = []
        errors: List[str] = []

        source_path = Path(source_document_path) if source_document_path else None
        if source_path and source_path.suffix.lower() == ".pdf":
            try:
                downloads.append(
                    self._export_pdf_overlay_result(
                        source_pdf_path=source_path,
                        original_name=original_name,
                        profile=profile,
                        autofill_result=autofill_result,
                    )
                )
            except ResultExportError as error:
                errors.append(str(error))
            if downloads:
                return downloads

        if source_path and source_path.suffix.lower() == ".docx":
            try:
                downloads.append(
                    self._export_docx_template_result(
                        source_docx_path=source_path,
                        original_name=original_name,
                        profile=profile,
                        autofill_result=autofill_result,
                    )
                )
            except ResultExportError as error:
                errors.append(str(error))
            if downloads:
                return downloads
            raise ResultExportError(" / ".join(errors) or "원본 DOCX 양식 채우기에 실패했습니다.")

        try:
            downloads.append(
                self._export_docx_result(
                    original_name=original_name,
                    profile=profile,
                    autofill_result=autofill_result,
                )
            )
        except ResultExportError as error:
            errors.append(str(error))

        if not downloads:
            raise ResultExportError(" / ".join(errors) or "결과 파일 생성에 실패했습니다.")
        return downloads

    def _export_docx_template_result(
        self,
        *,
        source_docx_path: Path,
        original_name: str,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> Dict[str, str]:
        if Document is None:
            raise ResultExportError("DOCX 생성에 필요한 python-docx를 사용할 수 없습니다.")

        stem = Path(original_name).stem or "document"
        safe_stem = self._safe_stem(stem)
        stored_name = f"{uuid.uuid4().hex}_{safe_stem}_filled_template.docx"
        output_path = self.output_dir / stored_name
        self._fill_original_docx(
            source_docx_path=source_docx_path,
            output_path=output_path,
            profile=profile,
            autofill_result=autofill_result,
        )
        return {
            "label": "원본 DOCX 채워진 결과 다운로드",
            "format": "docx",
            "stored_name": stored_name,
            "download_name": f"filled_{safe_stem}.docx",
            "download_url": f"/api/downloads/{stored_name}",
        }

    def _export_docx_result(
        self,
        *,
        original_name: str,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> Dict[str, str]:
        if Document is None or WD_ALIGN_PARAGRAPH is None or Pt is None:
            raise ResultExportError("DOCX 생성에 필요한 python-docx를 사용할 수 없습니다.")

        stem = Path(original_name).stem or "document"
        safe_stem = self._safe_stem(stem)
        stored_name = f"{uuid.uuid4().hex}_{safe_stem}_filled.docx"
        output_path = self.output_dir / stored_name
        self._build_docx(output_path, original_name, profile, autofill_result)
        return {
            "label": "채워진 DOCX 다운로드",
            "format": "docx",
            "stored_name": stored_name,
            "download_name": f"filled_{safe_stem}.docx",
            "download_url": f"/api/downloads/{stored_name}",
        }

    def _export_pdf_overlay_result(
        self,
        *,
        source_pdf_path: Path,
        original_name: str,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> Dict[str, str]:
        if PdfReader is None or PdfWriter is None or canvas is None or pdfmetrics is None:
            raise ResultExportError("PDF 결과 생성에 필요한 pypdf/reportlab를 사용할 수 없습니다.")

        reader = PdfReader(str(source_pdf_path))
        writer = self._overlay_pdf_answers(reader, profile, autofill_result)

        stem = Path(original_name).stem or "document"
        safe_stem = self._safe_stem(stem)
        stored_name = f"{uuid.uuid4().hex}_{safe_stem}_filled.pdf"
        output_path = self.output_dir / stored_name
        with output_path.open("wb") as handle:
            writer.write(handle)

        return {
            "label": "편집 가능한 원본 PDF 다운로드",
            "format": "pdf",
            "stored_name": stored_name,
            "download_name": f"filled_{safe_stem}.pdf",
            "download_url": f"/api/downloads/{stored_name}",
        }

    def _overlay_pdf_answers(self, reader, profile: Dict[str, Any], autofill_result: Dict[str, Any]):
        items_by_page = [self._extract_page_items(page) for page in reader.pages]
        page_sizes = [
            (float(page.mediabox.width), float(page.mediabox.height))
            for page in reader.pages
        ]
        placements_by_page = self._build_placements(items_by_page, page_sizes, profile, autofill_result)

        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        self._add_editable_pdf_fields(writer, placements_by_page)
        return writer

    def _add_editable_pdf_fields(self, writer, placements_by_page) -> None:
        if (
            ArrayObject is None
            or BooleanObject is None
            or DictionaryObject is None
            or NameObject is None
            or TextStringObject is None
        ):
            raise ResultExportError("PDF 편집 필드 생성에 필요한 pypdf 객체를 사용할 수 없습니다.")

        acroform = writer._root_object.get(NameObject("/AcroForm"))
        if acroform is None:
            acroform = DictionaryObject()
            writer._root_object[NameObject("/AcroForm")] = acroform
        else:
            acroform = acroform.get_object()

        if NameObject("/Fields") not in acroform:
            acroform[NameObject("/Fields")] = ArrayObject()
        fields = acroform[NameObject("/Fields")]

        acroform[NameObject("/NeedAppearances")] = BooleanObject(True)
        acroform[NameObject("/DA")] = TextStringObject("/Helv 9 Tf 0.08 0.20 0.45 rg")
        acroform[NameObject("/DR")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {
                        NameObject("/Helv"): DictionaryObject(
                            {
                                NameObject("/Type"): NameObject("/Font"),
                                NameObject("/Subtype"): NameObject("/Type1"),
                                NameObject("/BaseFont"): NameObject("/Helvetica"),
                            }
                        )
                    }
                )
            }
        )

        field_index = 1
        for page_index, placements in enumerate(placements_by_page):
            for placement in placements:
                annotation = self._build_pdf_text_field_annotation(placement, field_index)
                inserted = writer.add_annotation(page_index, annotation)
                fields.append(inserted.indirect_reference)
                field_index += 1

    def _build_pdf_text_field_annotation(self, placement: Dict[str, Any], field_index: int) -> Dict[str, Any]:
        font_size = float(placement["font_size"])
        max_width = float(placement["max_width"])
        max_height = float(placement["max_height"])
        x = float(placement["x"])
        top = float(placement["y"]) + font_size + 2.0
        bottom = max(0.0, top - max(14.0, max_height))
        value = str(placement["text"]).strip()

        return {
            "/Type": "/Annot",
            "/Subtype": "/Widget",
            "/FT": "/Tx",
            "/T": f"autofill_answer_{field_index}",
            "/V": value,
            "/DV": value,
            "/Rect": [x, bottom, x + max_width, top],
            "/F": 4,
            "/Ff": 4096,
            "/DA": f"/Helv {font_size:.1f} Tf 0.08 0.20 0.45 rg",
            "/BS": {
                "/W": 0,
            },
            "/Q": 0,
        }

    def _build_placements(
        self,
        items_by_page,
        page_sizes,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ):
        placements_by_page: List[List[Dict[str, Any]]] = [[] for _ in page_sizes]
        self._append_pdf_profile_placements(placements_by_page, items_by_page, page_sizes, profile)

        answers = [
            item for item in autofill_result.get("answers", [])
            if str(item.get("answer", "")).strip()
            and not self._is_pdf_non_answer_item(item)
        ]

        for answer in answers:
            match = self._match_answer_anchor(items_by_page, str(answer.get("question", "")))
            if match is None:
                continue
            page_index, item_index, score = match
            placement = self._build_answer_placement(
                items_by_page=items_by_page,
                page_sizes=page_sizes,
                page_index=page_index,
                item_index=item_index,
                answer=answer,
                score=score,
            )
            if placement:
                placements_by_page[placement["page_index"]].append(placement)

        return placements_by_page

    def _append_pdf_profile_placements(self, placements_by_page, items_by_page, page_sizes, profile: Dict[str, Any]) -> None:
        label_values = self._build_pdf_label_values(profile)
        if not label_values:
            return

        used_labels = set()
        for page_index, items in enumerate(items_by_page):
            for item_index, item in enumerate(items):
                normalized_label = self._normalize_label(item["text"])
                if not normalized_label or normalized_label in used_labels:
                    continue
                value = self._lookup_label_value(label_values, normalized_label)
                if not value:
                    continue

                placement = self._build_pdf_label_placement(
                    items=items,
                    page_sizes=page_sizes,
                    page_index=page_index,
                    item=item,
                    value=value,
                )
                if placement:
                    placements_by_page[placement["page_index"]].append(placement)
                    used_labels.add(normalized_label)

    def _build_pdf_label_placement(
        self,
        *,
        items: List[Dict[str, Any]],
        page_sizes,
        page_index: int,
        item: Dict[str, Any],
        value: str,
    ) -> Optional[Dict[str, Any]]:
        page_width, _ = page_sizes[page_index]
        label_width = self._measure_text(item["text"], max(item["font_size"], 9.0))
        same_row_right = [
            other["x"] for other in items
            if abs(other["y"] - item["y"]) < 4 and other["x"] > item["x"] + 40
        ]

        if same_row_right:
            next_label_x = min(same_row_right)
            start_x = item["x"] + max(54.0, label_width + 14.0)
            right_x = next_label_x - 10.0
        elif item["x"] < 180:
            start_x = item["x"] + max(130.0, label_width + 18.0)
            right_x = page_width - 55.0
        else:
            start_x = item["x"] + max(54.0, label_width + 14.0)
            right_x = page_width - 55.0

        if right_x - start_x < 42.0:
            return None

        return {
            "page_index": page_index,
            "x": start_x,
            "y": item["y"] - 1.0,
            "max_width": right_x - start_x,
            "max_height": 14.0,
            "font_size": 8.4,
            "text": value,
        }

    def _build_pdf_label_values(self, profile: Dict[str, Any]) -> Dict[str, str]:
        education = profile.get("학력") if isinstance(profile.get("학력"), dict) else {}
        profile_values = {
            "이름": str(profile.get("이름", "")).strip(),
            "생년월일": str(profile.get("생년월일", "")).strip(),
            "나이": str(profile.get("나이", "")).strip(),
            "이메일": str(profile.get("이메일", "")).strip(),
            "전화번호": str(profile.get("전화번호", "")).strip(),
            "주소": str(profile.get("주소", "")).strip(),
            "희망직무": str(profile.get("희망직무", "")).strip(),
            "지원직무": str(profile.get("희망직무", "")).strip(),
            "보유기술스택": ", ".join(profile.get("기술스택", []) or []),
            "사용가능언어": ", ".join(profile.get("기술스택", []) or []),
            "보유자격증": ", ".join(profile.get("자격증", []) or []),
            "관련자격증": ", ".join(profile.get("자격증", []) or []),
            "학교명": str(education.get("대학교", "")).strip(),
            "대학교": str(education.get("대학교", "")).strip(),
            "전공": str(education.get("전공", "")).strip(),
            "학점": str(education.get("학점", "")).strip(),
            "졸업예정년도": str(education.get("졸업년도", "")).strip(),
            "졸업년도": str(education.get("졸업년도", "")).strip(),
            "총경력": str(profile.get("경력", "")).strip(),
        }

        label_values: Dict[str, str] = {}
        for key, value in profile_values.items():
            normalized = self._normalize_label(key)
            if value and normalized:
                label_values[normalized] = value
        return label_values

    def _is_pdf_non_answer_item(self, item: Dict[str, Any]) -> bool:
        question = self._normalize_label(str(item.get("question", "")))
        section = self._normalize_label(str(item.get("section", "")))
        return "추가확인필요" in question or "추가확인필요" in section

    def _match_answer_anchor(self, items_by_page, question: str):
        question_norm = self._normalize_text(question)
        if not question_norm:
            return None

        best = None
        best_score = 0.0
        for page_index, items in enumerate(items_by_page):
            for item_index, item in enumerate(items):
                text = item["text"]
                if self._is_overlay_noise(text):
                    continue
                score = self._similarity(question_norm, self._normalize_text(text))
                if score > best_score:
                    best_score = score
                    best = (page_index, item_index, score)

        if best is None:
            return None
        threshold = 0.92 if len(question_norm) <= 4 else 0.48
        if best_score < threshold:
            return None
        return best

    def _build_answer_placement(
        self,
        *,
        items_by_page,
        page_sizes,
        page_index: int,
        item_index: int,
        answer: Dict[str, Any],
        score: float,
    ) -> Optional[Dict[str, Any]]:
        items = items_by_page[page_index]
        item = items[item_index]
        text = str(answer.get("answer", "")).strip()
        if not text:
            return None

        block = self._find_paragraph_block(items_by_page, page_sizes, page_index, item)
        if block is not None:
            return {
                "page_index": block["page_index"],
                "x": block["x"],
                "y": block["y"],
                "max_width": block["max_width"],
                "max_height": block["max_height"],
                "font_size": 9.0,
                "text": text,
            }

        page_width, _ = page_sizes[page_index]
        next_same_row_x = min(
            [other["x"] for other in items if abs(other["y"] - item["y"]) < 4 and other["x"] > item["x"] + 4] or [page_width - 40]
        )
        label_width = self._measure_text(item["text"], max(item["font_size"], 9.0))
        start_x = min(item["x"] + max(44.0, label_width + 10.0), next_same_row_x - 70.0)
        if start_x < item["x"] + 44.0:
            start_x = item["x"] + 44.0
        max_width = max(70.0, next_same_row_x - start_x - 12.0)

        gap_to_next = self._gap_to_next_label(items, item_index)
        if len(text) > 24 and gap_to_next >= 26:
            return {
                "page_index": page_index,
                "x": item["x"],
                "y": item["y"] - 14.0,
                "max_width": page_width - item["x"] - 55.0,
                "max_height": max(16.0, gap_to_next - 8.0),
                "font_size": 8.6,
                "text": text,
            }

        if len(text) > 18 and max_width < 120 and gap_to_next >= 18:
            return {
                "page_index": page_index,
                "x": start_x,
                "y": item["y"] - 2.0,
                "max_width": max_width,
                "max_height": 18.0,
                "font_size": 8.5,
                "text": text,
            }

        return {
            "page_index": page_index,
            "x": start_x,
            "y": item["y"] - 1.0,
            "max_width": max_width,
            "max_height": 14.0,
            "font_size": 9.0 if len(text) <= 20 else 8.2,
            "text": text,
        }

    def _find_paragraph_block(self, items_by_page, page_sizes, page_index: int, anchor_item: Dict[str, Any]):
        items = items_by_page[page_index]
        underscore_lines = [
            item for item in items
            if self._is_long_underscore_line(item["text"]) and 0 < anchor_item["y"] - item["y"] < 120
        ]
        if underscore_lines:
            top_y = max(item["y"] for item in underscore_lines)
            bottom_y = min(item["y"] for item in underscore_lines)
            page_width, _ = page_sizes[page_index]
            return {
                "page_index": page_index,
                "x": max(58.0, min(item["x"] for item in underscore_lines) + 4.0),
                "y": top_y + 2.0,
                "max_width": page_width - 110.0,
                "max_height": max(24.0, top_y - bottom_y + 12.0),
            }

        if anchor_item["y"] < 110 and page_index + 1 < len(items_by_page):
            next_items = items_by_page[page_index + 1]
            underscore_lines = [
                item for item in next_items
                if self._is_long_underscore_line(item["text"]) and item["y"] > 680
            ]
            if underscore_lines:
                top_y = max(item["y"] for item in underscore_lines)
                bottom_y = min(item["y"] for item in underscore_lines)
                page_width, _ = page_sizes[page_index + 1]
                return {
                    "page_index": page_index + 1,
                    "x": max(58.0, min(item["x"] for item in underscore_lines) + 4.0),
                    "y": top_y + 2.0,
                    "max_width": page_width - 110.0,
                    "max_height": max(24.0, top_y - bottom_y + 12.0),
                }
        return None

    def _build_overlay_pdf(self, page_sizes, placements_by_page):
        buffer = BytesIO()
        first_width, first_height = page_sizes[0]
        pdf = canvas.Canvas(buffer, pagesize=(first_width, first_height))

        for page_index, (page_width, page_height) in enumerate(page_sizes):
            if page_index > 0:
                pdf.setPageSize((page_width, page_height))
            for placement in placements_by_page[page_index]:
                self._draw_wrapped_text(pdf, placement)
            pdf.showPage()

        pdf.save()
        return buffer.getvalue()

    def _draw_wrapped_text(self, pdf, placement: Dict[str, Any]) -> None:
        font_size = float(placement["font_size"])
        max_width = float(placement["max_width"])
        max_height = float(placement["max_height"])
        line_height = font_size + 2.4
        max_lines = max(1, int(max_height / line_height))
        lines = self._wrap_text(
            text=str(placement["text"]),
            font_size=font_size,
            max_width=max_width,
            max_lines=max_lines,
        )
        pdf.setFillColorRGB(0.08, 0.20, 0.45)
        pdf.setFont(self.pdf_font_name, font_size)
        current_y = float(placement["y"])
        for line in lines:
            pdf.drawString(float(placement["x"]), current_y, line)
            current_y -= line_height

    def _wrap_text(self, *, text: str, font_size: float, max_width: float, max_lines: int) -> List[str]:
        text = text.replace("\r", "").strip()
        if not text:
            return [""]

        lines: List[str] = []
        for paragraph in text.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            current = ""
            for char in paragraph:
                candidate = current + char
                if self._measure_text(candidate, font_size) <= max_width:
                    current = candidate
                    continue
                if current:
                    lines.append(current.rstrip())
                current = char
                if len(lines) >= max_lines:
                    return self._truncate_lines(lines, font_size, max_width, max_lines)
            if current:
                lines.append(current.rstrip())
                if len(lines) >= max_lines:
                    return self._truncate_lines(lines, font_size, max_width, max_lines)
        return self._truncate_lines(lines, font_size, max_width, max_lines)

    def _truncate_lines(self, lines: List[str], font_size: float, max_width: float, max_lines: int) -> List[str]:
        if len(lines) <= max_lines:
            return lines
        trimmed = lines[:max_lines]
        ellipsis = "..."
        last = trimmed[-1]
        while last and self._measure_text(last + ellipsis, font_size) > max_width:
            last = last[:-1]
        trimmed[-1] = (last + ellipsis).strip() if last else ellipsis
        return trimmed

    def _measure_text(self, text: str, font_size: float) -> float:
        if pdfmetrics is None:
            return len(text) * font_size
        return pdfmetrics.stringWidth(text, self.pdf_font_name, font_size)

    def _extract_page_items(self, page) -> List[Dict[str, Any]]:
        items: List[Dict[str, Any]] = []

        def visitor_text(text, cm, tm, font_dict, font_size):
            content = str(text).strip()
            if not content:
                return
            items.append(
                {
                    "x": float(tm[4]),
                    "y": float(tm[5]),
                    "font_size": float(font_size),
                    "text": content,
                }
            )

        try:
            page.extract_text(visitor_text=visitor_text)
        except TypeError:
            page.extract_text()
        return items

    def _gap_to_next_label(self, items: List[Dict[str, Any]], item_index: int) -> float:
        current = items[item_index]
        candidates = [
            item["y"] for idx, item in enumerate(items)
            if idx != item_index
            and abs(item["x"] - current["x"]) < 40
            and 0 < current["y"] - item["y"] < 120
            and not self._is_overlay_noise(item["text"])
        ]
        if not candidates:
            return 22.0
        return current["y"] - max(candidates)

    def _similarity(self, left: str, right: str) -> float:
        if not left or not right:
            return 0.0
        if left == right:
            return 2.0
        if len(left) <= 4 or len(right) <= 4:
            return 1.2 if left == right else 0.0
        ratio = SequenceMatcher(None, left, right).ratio()
        if left in right or right in left:
            ratio += 0.35
        return ratio

    def _normalize_text(self, value: str) -> str:
        lowered = str(value).lower()
        return re.sub(r"[^0-9a-zA-Z가-힣]", "", lowered)

    def _is_overlay_noise(self, text: str) -> bool:
        return self._is_long_underscore_line(text) or text == "__" or text == "." or text == "/" or text == "(" or text == ")"

    def _is_long_underscore_line(self, text: str) -> bool:
        stripped = text.strip()
        return len(stripped) >= 20 and set(stripped) == {"_"}

    def _configure_pdf_font(self) -> str:
        default_font = "Helvetica"
        if pdfmetrics is None or TTFont is None:
            return default_font

        font_name = "PersonaPdfFont"
        try:
            pdfmetrics.getFont(font_name)
            return font_name
        except KeyError:
            pass

        font_candidates = [
            Path("/System/Library/Fonts/Supplemental/AppleGothic.ttf"),
            Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        ]
        for path in font_candidates:
            if not path.exists():
                continue
            try:
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
                return font_name
            except Exception:
                continue
        return default_font

    def _build_docx(
        self,
        output_path: Path,
        original_name: str,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> None:
        document = Document()

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("자동입력 결과 문서")
        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(
            f"원본 문서: {original_name} | 생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        subtitle_run.font.size = Pt(10)

        document.add_paragraph("")

        info = document.add_paragraph()
        info_run = info.add_run(
            f"프로필 기준: {profile.get('이름', '')} / {profile.get('희망직무', '')}"
        )
        info_run.font.size = Pt(11)

        summary = document.add_paragraph()
        summary_run = summary.add_run(
            autofill_result.get('summary', '문서 요약이 생성되지 않았습니다.')
        )
        summary_run.font.size = Pt(11)

        document.add_paragraph("")

        current_section = None
        for answer in autofill_result.get('answers', []):
            section = (answer.get('section') or '기타').strip()
            if section != current_section:
                current_section = section
                section_p = document.add_paragraph()
                section_run = section_p.add_run(section)
                section_run.bold = True
                section_run.font.size = Pt(14)

            question_p = document.add_paragraph()
            question_run = question_p.add_run(answer.get('question', ''))
            question_run.bold = True
            question_run.font.size = Pt(11)

            answer_p = document.add_paragraph(answer.get('answer', ''))
            if answer_p.runs:
                answer_p.runs[0].font.size = Pt(11)

            reason = (answer.get('reason') or '').strip()
            if reason:
                reason_p = document.add_paragraph(f"사유: {reason}")
                if reason_p.runs:
                    reason_p.runs[0].font.size = Pt(9)

            source = answer.get('source') or []
            if source:
                source_p = document.add_paragraph(f"근거: {', '.join(source)}")
                if source_p.runs:
                    source_p.runs[0].font.size = Pt(9)

            document.add_paragraph("")

        missing_information = autofill_result.get('missing_information') or []
        if missing_information:
            missing_title = document.add_paragraph()
            missing_run = missing_title.add_run("추가 확인 필요")
            missing_run.bold = True
            missing_run.font.size = Pt(14)

            for item in missing_information:
                line = f"- {item.get('question', '')}: {item.get('reason', '')}"
                paragraph = document.add_paragraph(line)
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(10)

        document.save(str(output_path))

    def _fill_original_docx(
        self,
        *,
        source_docx_path: Path,
        output_path: Path,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> None:
        if Document is None:
            raise ResultExportError("DOCX 템플릿 채우기에 필요한 python-docx를 사용할 수 없습니다.")

        try:
            document = Document(str(source_docx_path))
        except Exception as error:
            raise ResultExportError(f"원본 DOCX를 열 수 없습니다: {error}") from error

        label_values = self._build_docx_label_values(profile, autofill_result)
        self._fill_docx_label_tables(document, label_values)
        self._fill_docx_prompt_blocks(document, autofill_result)
        self._fill_docx_inline_placeholders(document, label_values)
        document.save(str(output_path))

    def _build_docx_label_values(
        self,
        profile: Dict[str, Any],
        autofill_result: Dict[str, Any],
    ) -> Dict[str, str]:
        answers = autofill_result.get("answers") or []
        answer_values: Dict[str, str] = {}
        for item in answers:
            if not isinstance(item, dict):
                continue
            answer_text = str(item.get("answer", "")).strip()
            if not answer_text:
                continue
            for key in [
                str(item.get("question", "")).strip(),
                str(item.get("field_key", "")).strip(),
            ]:
                normalized = self._normalize_label(key)
                if normalized and normalized not in answer_values:
                    answer_values[normalized] = answer_text

        education = profile.get("학력") if isinstance(profile.get("학력"), dict) else {}
        profile_values = {
            "이름": str(profile.get("이름", "")).strip(),
            "생년월일": str(profile.get("생년월일", "")).strip(),
            "나이": str(profile.get("나이", "")).strip(),
            "이메일": str(profile.get("이메일", "")).strip(),
            "전화번호": str(profile.get("전화번호", "")).strip(),
            "주소": str(profile.get("주소", "")).strip(),
            "희망직무": str(profile.get("희망직무", "")).strip(),
            "지원직무": str(profile.get("희망직무", "")).strip(),
            "보유기술스택": ", ".join(profile.get("기술스택", []) or []),
            "사용가능언어": ", ".join(profile.get("기술스택", []) or []),
            "관련자격증": ", ".join(profile.get("자격증", []) or []),
            "보유자격증": ", ".join(profile.get("자격증", []) or []),
            "학교명": str(education.get("대학교", "")).strip(),
            "대학교": str(education.get("대학교", "")).strip(),
            "전공": str(education.get("전공", "")).strip(),
            "학점": str(education.get("학점", "")).strip(),
            "졸업예정년도": str(education.get("졸업년도", "")).strip(),
            "졸업년도": str(education.get("졸업년도", "")).strip(),
            "총경력": str(profile.get("경력", "")).strip(),
            "경력": str(profile.get("경력", "")).strip(),
            "추가메모": str(profile.get("추가메모", "")).strip(),
        }

        label_values: Dict[str, str] = {}
        for key, value in profile_values.items():
            normalized = self._normalize_label(key)
            if value and normalized:
                label_values[normalized] = value

        label_values.update(answer_values)
        return label_values

    def _fill_docx_label_tables(self, document, label_values: Dict[str, str]) -> None:
        for table in document.tables:
            for row in table.rows:
                cells = row.cells
                for index, cell in enumerate(cells[:-1]):
                    label = self._normalize_label(cell.text)
                    if not label:
                        continue
                    value = self._lookup_label_value(label_values, label)
                    if not value:
                        continue
                    next_cell = cells[index + 1]
                    if next_cell.text.strip():
                        continue
                    self._set_docx_cell_text(next_cell, value)

    def _fill_docx_prompt_blocks(self, document, autofill_result: Dict[str, Any]) -> None:
        prompt_answers = self._build_prompt_answer_map(autofill_result)
        ordered_prompt_answers = self._build_ordered_prompt_answers(autofill_result)
        if not prompt_answers:
            return

        table_index = 0
        used_prompt_keys: set[str] = set()
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue
            answer_text, matched_key = self._match_prompt_answer(
                prompt_answers,
                paragraph_text,
                used_prompt_keys,
            )
            if not answer_text and self._looks_like_prompt_paragraph(paragraph_text):
                fallback = self._next_unused_prompt_answer(ordered_prompt_answers, used_prompt_keys)
                if fallback:
                    matched_key, answer_text = fallback
            if not answer_text:
                continue
            while table_index < len(document.tables):
                candidate = document.tables[table_index]
                table_index += 1
                if not self._is_docx_blank_answer_block(candidate):
                    continue
                target_cell = candidate.rows[0].cells[0]
                if self._is_meaningful_text(target_cell.text):
                    continue
                self._set_docx_cell_text(target_cell, answer_text)
                if matched_key:
                    used_prompt_keys.add(matched_key)
                break

    def _fill_docx_inline_placeholders(self, document, label_values: Dict[str, str]) -> None:
        for paragraph in document.paragraphs:
            updated = self._replace_inline_placeholder_text(paragraph.text, label_values)
            if updated != paragraph.text:
                self._set_docx_paragraph_text(paragraph, updated)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        updated = self._replace_inline_placeholder_text(paragraph.text, label_values)
                        if updated != paragraph.text:
                            self._set_docx_paragraph_text(paragraph, updated)

    def _build_prompt_answer_map(self, autofill_result: Dict[str, Any]) -> Dict[str, str]:
        prompt_answers: Dict[str, str] = {}
        for item in autofill_result.get("answers", []):
            if not isinstance(item, dict):
                continue
            question = str(item.get("question", "")).strip()
            answer = str(item.get("answer", "")).strip()
            if not question or not answer:
                continue
            normalized = self._normalize_label(question)
            if normalized:
                prompt_answers[normalized] = answer
        return prompt_answers

    def _build_ordered_prompt_answers(self, autofill_result: Dict[str, Any]) -> List[tuple[str, str]]:
        ordered: List[tuple[str, str]] = []
        for item in autofill_result.get("answers", []):
            if not isinstance(item, dict):
                continue
            question = str(item.get("question", "")).strip()
            answer = str(item.get("answer", "")).strip()
            normalized = self._normalize_label(question)
            if not normalized or not answer:
                continue
            if not self._looks_like_prompt_paragraph(question):
                continue
            ordered.append((normalized, answer))
        return ordered

    def _match_prompt_answer(
        self,
        prompt_answers: Dict[str, str],
        paragraph_text: str,
        used_prompt_keys: set[str],
    ) -> tuple[str, str]:
        normalized_paragraph = self._normalize_label(paragraph_text)
        if not normalized_paragraph:
            return "", ""

        direct = prompt_answers.get(normalized_paragraph)
        if direct and normalized_paragraph not in used_prompt_keys:
            return direct, normalized_paragraph

        best_answer = ""
        best_key = ""
        best_score = 0.0
        for key, answer in prompt_answers.items():
            if key in used_prompt_keys:
                continue
            score = self._similarity(normalized_paragraph, key)
            if score > best_score:
                best_score = score
                best_answer = answer
                best_key = key
        if best_score >= 0.55:
            return best_answer, best_key
        return "", ""

    def _next_unused_prompt_answer(
        self,
        ordered_prompt_answers: List[tuple[str, str]],
        used_prompt_keys: set[str],
    ) -> tuple[str, str] | None:
        for key, answer in ordered_prompt_answers:
            if key in used_prompt_keys:
                continue
            return key, answer
        return None

    def _looks_like_prompt_paragraph(self, text: str) -> bool:
        stripped = str(text or "").strip()
        if not stripped:
            return False
        if "작성해 주세요" in stripped or "작성해주세요" in stripped:
            return True
        if stripped.endswith("?"):
            return True
        return bool(re.match(r"^\d+\.", stripped))

    def _replace_inline_placeholder_text(self, text: str, label_values: Dict[str, str]) -> str:
        if ":" not in text or "_" not in text:
            return text
        left, _ = text.split(":", 1)
        label = self._normalize_label(left)
        value = self._lookup_label_value(label_values, label)
        if not value:
            return text
        return f"{left}: {value}"

    def _lookup_label_value(self, label_values: Dict[str, str], label: str) -> str:
        direct = label_values.get(label)
        if direct:
            return direct

        aliases = {
            "지원직무": ["희망직무"],
            "희망직무": ["지원직무"],
            "학교명": ["대학교"],
            "대학교": ["학교명"],
            "관련자격증": ["보유자격증"],
            "보유자격증": ["관련자격증"],
            "졸업예정년도": ["졸업년도"],
            "졸업년도": ["졸업예정년도"],
            "보유기술스택": ["기술스택"],
            "기술스택": ["보유기술스택"],
        }
        for alias in aliases.get(label, []):
            value = label_values.get(alias)
            if value:
                return value
        return ""

    def _normalize_label(self, value: str) -> str:
        return re.sub(r"[^0-9a-zA-Z가-힣]", "", str(value or "").strip().lower())

    def _is_docx_blank_answer_block(self, table) -> bool:
        if len(table.rows) != 1 or len(table.columns) != 1:
            return False
        text = table.rows[0].cells[0].text
        return not self._is_meaningful_text(text)

    def _is_meaningful_text(self, text: str) -> bool:
        stripped = str(text or "").strip()
        if not stripped:
            return False
        visible = [char for char in stripped if not char.isspace()]
        if not visible:
            return False
        punctuation_ratio = sum(1 for char in visible if char in "_-=[]()|.:,/") / len(visible)
        return punctuation_ratio < 0.7

    def _set_docx_cell_text(self, cell, value: str) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        if Pt is not None:
            run.font.size = Pt(10.5)

    def _set_docx_paragraph_text(self, paragraph, value: str) -> None:
        paragraph.text = ""
        run = paragraph.add_run(value)
        if Pt is not None:
            run.font.size = Pt(10.5)

    def _safe_stem(self, stem: str) -> str:
        cleaned = []
        for char in stem:
            if char.isalnum() or char in {'-', '_'}:
                cleaned.append(char)
            elif char in {' ', '.'}:
                cleaned.append('_')
        value = ''.join(cleaned).strip('_')
        return value or 'document'
